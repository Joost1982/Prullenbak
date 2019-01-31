# Script om tabellen in een Word bestand op alfabetische volgorde te zetten
#
# joostvdhorst@gmail.com

from docx import Document
from docx.shared import Pt # voor lettergrootte
import os


# CONFIG

pad = os.getcwd()
bestand = 'test.docx' # tabel met 3 regels
bestand2 = 'test_gesorteerd.docx'


def leesDoc():
    wordDoc = Document(pad+'/'+bestand)
    tabellen = []

    for table in wordDoc.tables:
        tabellen.append(table)

    regel1 = []
    regel2 = []
    regel3 = []

    for i in range(0,len(tabellen)):
        regel1.append(wordDoc.tables[i].rows[0].cells[1].text)
        regel2.append(wordDoc.tables[i].rows[1].cells[1].text)
        regel3.append(wordDoc.tables[i].rows[2].cells[1].text)

    return wordDoc, regel1, regel2, regel3


def sorteer():
    totaal = zip(regel1, regel2, regel3)
    totaal.sort()
    regel1b = []
    regel2b = []
    regel3b = []
    regel1b, regel2b, regel3b = zip(*totaal)

    return regel1b, regel2b, regel3b


def schrijfDoc():
    style = wordDoc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.bold = True

    #regel 1
    for i in range(0, len(regel1b)):
        wordDoc.tables[i].rows[0].cells[1].text = regel1b[i]

    #regel 2
    for i in range(0, len(regel2b)):
        wordDoc.tables[i].rows[1].cells[1].text = regel2b[i]

    #regel 3
    for i in range(0, len(regel3b)):
        wordDoc.tables[i].rows[2].cells[1].text = regel3b[i]

    # Document wegschrijven onder een ander naam
    wordDoc.save(pad+'/'+bestand2)


# Main
print '%s uitlezen en wegschrijven naar:' % (bestand)
print pad+'\\'+bestand2
print

wordDoc, regel1, regel2, regel3 = leesDoc()

regel1b, regel2b, regel3b = sorteer()

schrijfDoc()

print 'Klaar.'
print

raw_input('Druk op [ENTER] om af te sluiten.')
